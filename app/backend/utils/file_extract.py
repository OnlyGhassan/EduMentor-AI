import io
import PyPDF2
import docx
from fastapi import UploadFile

def extract_text_from_file(file: UploadFile) -> str:
    try:
        filename = file.filename.lower()
        file.file.seek(0)  # Ensure reading from start

        if filename.endswith(".pdf"):
            reader = PyPDF2.PdfReader(file.file)
            return "\n".join([page.extract_text() or "" for page in reader.pages])

        elif filename.endswith(".docx"):
            bio = io.BytesIO(file.file.read())
            doc = docx.Document(bio)
            return "\n".join([p.text for p in doc.paragraphs])

        elif filename.endswith(".txt"):
            content = file.file.read().decode("utf-8", errors="ignore")
            return f"Beginning of a single file {{ {content} }} end of a single file"

        else:
            return ""
    except Exception as e:
        print(f"Error: {e}")
        return ""


# def extract_text_from_file(file: UploadFile) -> str:
#     try:
#         filename = (file.filename or "").lower()
#         file.file.seek(0)

#         if filename.endswith(".pdf"):
#             reader = PyPDF2.PdfReader(file.file)
#             return "\n".join([page.extract_text() or "" for page in reader.pages])

#         if filename.endswith(".docx"):
#             bio = io.BytesIO(file.file.read())
#             d = docx.Document(bio)
#             return "\n".join(p.text for p in d.paragraphs)

#         if filename.endswith(".txt"):
#             return file.file.read().decode("utf-8", errors="ignore")

#         return ""
#     except Exception:
#         return ""
